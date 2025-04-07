import re
from datetime import datetime
import docx
import pandas as pd
from unidecode import unidecode
from cryptography.fernet import Fernet

# PARAMETERS
REPORT_NUMBER = {"table": 0, "cell": (0, 2)}
VERSION = {"table": 0, "cell": (0, 4)}
F4E_REFERENCE = {"table": 0, "cell": (1, 2)}
CUSTOMER_REF = {"table": 0, "cell": (1, 4)}
DMS_CELL = {"table": 0, "cell": (2, 2)}
KOM_DATE = {"table": 1, "cell": (1, 1)}
H_IN_PERIOD_CELL = {"table": 1, "cell": (1, 3)}
PERIODS = [
    {"table": 2, "cell": (0, 1), "section": "2.2"},
    {"table": 3, "cell": (0, 1), "section": "2.3"},
    {"table": 4, "cell": (0, 1), "section": "2.4"},
]
AUTHOR_NAME = {"table": 6, "cell": (1, 0)}
DATE_AUTHOR = {"table": 6, "cell": (2, 0)}
DATE_APPROVAL = {"table": 6, "cell": (2, 1)}
NEW_MILESTONE = {"table": 2, "cell": (1, 1)}
CURRENT_MILESTONE = {"table": 3, "cell": (1, 1)}
MILESTONE_TO_COPY = {"table": 4, "cell": (1, 1)}
TOTAL_HOURS = {"table": 1, "cell": (1, 3)}
SECTION3 = {"table": 3, "cell": (1, 1)}
TO_HIGHLIGHT = [
    SECTION3,
    {"table": 5, "cell": (1, 1)},
    MILESTONE_TO_COPY,
]
KEY_ENCRYPTED = {"table": 6, "cell": (3, 0)}
KEY_ENCRYPTED_SIDE = {"table": 6, "cell": (3, 1)}

FOLDER = "test_data"
MONTH_NUMBER_TO_NAME = {
    1: "January",
    2: "February",
    3: "March",
    4: "April",
    5: "May",
    6: "June",
    7: "July",
    8: "August",
    9: "September",
    10: "October",
    11: "November",
    12: "December",
}

global hours_task_plan

def process_mr(mr_files, hours_task_plan):

    global results_df
    results_df = pd.DataFrame(columns=["Reference", "Name", "Error"])
    hours_task_plan = pd.read_excel(hours_task_plan, skiprows=3)
    list_employees = pd.read_excel(r"D:\DATA\ferrmar\Documents\04-ATG\automatic_monthly_check\webapp\Development\LIST OF EMPLOYEES.xlsx")
    list_employees = list_employees[list_employees["Contract status"] == "Active"]
    for report in mr_files:
        process_monthly(report, hours_task_plan, list_employees)

    return results_df


def diff_month(d1, d2):
    """
    Returns the difference of months between two dates

    Arguments:
        d1 (datetime)
        d2 (datetime)
    Returns:
        int
    """
    return (d1.year - d2.year) * 12 + d1.month - d2.month


class HeaderData:
    def __init__(self):
        self.report_number = None
        self.f4e_reference = None
        self.supplier_dms = None
        self.kom_date = None
        self.reported_hours = None
        self.version = None
        self.customer_ref = None

    @property
    def totally_filled(self):
        return (
            self.report_number
            and self.f4e_reference
            and self.supplier_dms
            and self.kom_date
            and self.reported_hours
            and self.version
            and self.customer_ref
        )

    def __str__(self):
        text = f"""
Report number: {self.report_number}
F4E reference: {self.f4e_reference}
Supplier DMS: {self.supplier_dms}
KOM Date: {self.kom_date}
Reported hours: {self.reported_hours}
Revision: {self.version}
Customer Ref.: {self.customer_ref}"""
        return text


class PersonData:
    def __init__(self, list_employees):
        self.df = list_employees
        self.contract = None
        self.kom = None
        self.dms = None
        self.name_monthly = None
        self.name_irs = None
        self.row_data = None
        self.customer_ref = None

    def select_row(self, name_report):
        self.row_data = self.df[self.df["Name Monthly/Mission"].astype(str).apply(unidecode) == unidecode(name_report)]
        self.define_data()
        return
    
    def define_data(self):
        self.contract = self.row_data["Specific Contract"].values[0]
        self.kom = self.row_data["KoM"].values[0]
        self.name_monthly = unidecode(self.row_data["Name Monthly/Mission"].values[0])
        self.name_irs = unidecode(self.row_data["Name IRS"].values[0])
        self.customer_ref = self.row_data["F4E Customer Ref"].values[0]
        return
    
    def get_dms(self, month, year):
        self.dms = self.row_data[f"{month} {year}"].values[0]
        return
    

class Hours:
    def __init__(self):
        self.emt_total = None
        self.emt_general = None
        self.emt_specific = None
        self.report_total = None
        self.report_general = None
        self.report_specific = None
        self.report_general_taskplan = None
        self.report_taskplan_dic = None

    def hours_extmytime(self, hours_task_plan, person_data):
        hours_emt = hours_task_plan[hours_task_plan["Full Name"].apply(unidecode).isin([person_data.name_irs, person_data.name_monthly])]
        if len(hours_emt) == 0:
            error_message = f"  The name '{person_data.name_irs}' could not be found in the ExtMyTime file and, consequently, the hours " \
                            f"couldn't be checked"
            print(error_message)
            results_df.loc[len(results_df)] = [person_data.contract, name_report, error_message[2:]]
            self.emt_total = 0
            self.emt_general = 0
            self.emt_specific = {}
            return
        self.emt_total = hours_emt["Total Working hours submitted"].sum()
        self.emt_general = hours_emt[hours_emt["Task Plan Description"].str.contains("General Activities")]["Total Working hours submitted"].values[0]
        hours_specific = hours_emt[~hours_emt["Task Plan Description"].str.contains("General Activities")]
        self.emt_specific = dict(zip(hours_specific["Task Plan Code"], hours_specific["Total Working hours submitted"]))
        return

    def hours_report(self, document, header_data):
        self.report_general = 0
        self.report_specific = 0
        self.report_general_taskplan = ""
        self.report_taskplan_dic = {}
        
        section = document.tables[CURRENT_MILESTONE["table"]].cell(*CURRENT_MILESTONE["cell"]).text
        while True:
            try:
                match = re.search(r"Task.+:\s*(\d+([,.]\d+)?)\s*hours?\s*", section)
                section = section[match.span()[1]:]
            except AttributeError:
                break
            hours_task = match.group(1).replace(",", ".")
            line = match.group(0)
            try:
                hours_task = float(hours_task)
                if "General Activities".casefold() in line.casefold():
                    self.report_general += hours_task
                    self.report_general_taskplan = line[line.find("(") + 1:line.find(")")]
                else:
                    self.report_specific += hours_task
                    self.report_taskplan_dic[int(line[line.find("(") + 1:line.find(")")])] = hours_task
            except ValueError:
                error_message = f"  The number of hours in the line '{line}' could not be transformed to a number." \
                                f"Check if it is written correctly. The hours of that task could not be processed and," \
                                f"probably, the total number of hour will be incorrect because of this."
                print(error_message)
                results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
            self.report_total = self.report_general + self.report_specific
        return

        """
        except ValueError:
            error_message = f"  The code of the General Activities Task Plan in the report does not match the valid " \
                            f"format: {hours.report_general_taskplan}. The number of hours in the Task Plan could not be compared " \
                            f"between ExtMyTime and the report."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
        """
        """
        except ValueError:
            error_message = f"  The code of the General Activities Task Plan in the report does not match the valid " \
                            f"format: {hours.report_general_taskplan}. The number of hours in the Task Plan could not be compared " \
                            f"between ExtMyTime and the report."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
        """
        
        """
        # THIS ERROR HAS TO BE TRANSLATED TO THE HOURS CLASS
        except ValueError:
            error_message = f"  The code of a Specific Task Plan in the report does not match the valid format: " \
                            f"{specific_task}. The number of hours in the Task Plan could not be compared between " \
                            f"ExtMyTime and the report."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
        """


def read_header(document):
    """
    Extracts the data from the header of the monthly report

    Parameters:
        document (str): string with the text of the report

    Returns
        HeaderData class
    """
    header_data = HeaderData()
    header_data.report_number = document.tables[REPORT_NUMBER["table"]].cell(*REPORT_NUMBER["cell"]).text.strip()
    header_data.version = document.tables[VERSION["table"]].cell(*VERSION["cell"]).text.strip()
    header_data.f4e_reference = document.tables[F4E_REFERENCE["table"]].cell(*F4E_REFERENCE["cell"]).text.strip()
    header_data.customer_ref = document.tables[CUSTOMER_REF["table"]].cell(*CUSTOMER_REF["cell"]).text.strip()
    header_data.supplier_dms = document.tables[DMS_CELL["table"]].cell(*DMS_CELL["cell"]).text.strip()
    header_data.kom_date = document.tables[KOM_DATE["table"]].cell(*KOM_DATE["cell"]).text.strip()
    hours = document.tables[H_IN_PERIOD_CELL["table"]].cell(*H_IN_PERIOD_CELL["cell"]).text.strip()
    to_be_replaced = [" ", "(", "*", ")"]
    for symbol in to_be_replaced:
        hours = hours.replace(symbol, "")
    header_data.reported_hours = hours
    
    if not header_data.totally_filled:
        error_message = f"  The monthly header was not successfully read {header_data}"
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
        raise ValueError(error_message)
    return header_data


def add_header_hours_to_list(header_data):
    """
    Adds in the total_hours_df the hours done by the person during the period, as per the information
    in the report header

    Parameters:
        header_data (class)
    """
    total_hours_df.loc[len(total_hours_df)] = [header_data.f4e_reference, name_report, header_data.reported_hours]
    return


def show_version_message(header_data):
    """
    Prints a messages with the Revision number of the document
    :param header_data(class)
    :return: None
    """
    print(f"  The revision number in the header is {header_data.version}. Make sure this is correct.")
    return


def get_names(filename, list_employees):
    """
    Initializes class name, reads the name in the title of the report and converts it to the other two formats.

    Parameters:
        list_names (dataframe): extracted from the file with all the names and their equivalents
        filename (str): file name of the monthly report file
    Returns:
        name (class)
    """
    name = Name()
    name.report = unidecode(re.match(r".+ Monthly Report (.+\s.+) #", filename).group(1))
    name.convert(list_employees)

    return name


def check_f4e_contract(code_from_filename, header_data, person_data):
    """Check if F4E reference is the same in the name of the report and inside the report"""
    if header_data.f4e_reference not in code_from_filename:
        error_message = f"  The F4E contract shown in the header ({header_data.f4e_reference}) does not match " \
                        f"the one of the Word filename ({code_from_filename})"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    if header_data.f4e_reference != person_data.contract:
        error_message = f"  The F4E contract shown in the header ({header_data.f4e_reference}) does not match " \
                        f"the one in the LIST OF EMPLOYEES file ({person_data.contract})"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


def check_supplier_dms(header_data, person_data):
    """Check if the report of the month, year and person is in the Excel file 'DMS Number Monthly Report.xlsx' and
    also if the DMS number is correspondent"""

    # Get month from report number in header
    month = re.match(r"#\d+_M(\d+)_\d+", header_data.report_number).group(1)  # e.g. "3"
    month = int(month)
    month = MONTH_NUMBER_TO_NAME[month]  # e.g. "March"
    # Get year from report number in header
    year = re.match(r"#\d+_M\d+_(\d+)", header_data.report_number).group(1)  # eg."2022"
    # Get name from filename
    """
    if name.irs is None:
        error_message = f"  The name '{name.report}' could not be found in the file with the list of names under the " \
                        f"column named 'Monthly/Mission request'. The DMS number could not be checked."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
        return
    """
    # Search DMS and compare to report
    person_data.get_dms(month, year)
    if not pd.isna(person_data.dms):
        #dms_code = dms.values[0]
        if person_data.dms != header_data.supplier_dms:
            error_message = f"  DMS from database ({person_data.dms}) does not match DMS from " \
                            f"header ({header_data.supplier_dms}). Check the DMS number and the month number in the " \
                            f"report."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    else:
        error_message = f"  The DMS reference could not be found for {month} {year} for {person_data.name_monhlty}. " \
                        f"Either the DMS is not in the file or any of these parameters is written incorrectly. " \
                        f"It could not be checked if the DMS number is correct."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


def check_report_number_against_kom_date(header_data):
    """Check if the report number is coherent with the months passed since the KoM"""
    # Get KOM date
    day, month, year = [int(x) for x in header_data.kom_date.split(r"/")]
    kom_date = datetime(year, month, day)
    # Get month, year and report number from header
    number, month, year = header_data.report_number.split("_")
    number, month, year = int(number[1:]), int(month[1:]), int(year)
    report_date = datetime(year, month, 1)
    # Check if difference of months is equal to report number
    months_passed = diff_month(report_date, kom_date) + 1
    if months_passed != number:
        error_message = f"  The F4E report number ({header_data.report_number}) and the KOM date " \
                        f"({header_data.kom_date}) have a difference of {months_passed} months"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


def check_customer_ref(header_data, person_data):

    if not pd.isna(person_data.customer_ref):
        if person_data.customer_ref != header_data.customer_ref:
            error_message = f"  The F4E Customer Reference in the report ({header_data.customer_ref}) is different" \
                            f"from the correct reference ({person_data.customer_ref})"
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    else:
        error_message = f"  Could not find a F4E Customer Reference in the LIST OF EMPLOYEES file."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return 


def check_hours_report_vs_header(header_data, hours):
    """Check if the number of total hours in the report is the same as the sum of the general activities and the
    specific tasks in the report"""
    
    # Check if sum of hours in the report is the same as in the header
    if not almost_equal(float(header_data.reported_hours), hours.report_total):
        error_message = f"  The sum of hours of the tasks ({hours.report_total}) is not the same as the one found " \
                        f"in the header ({header_data.reported_hours})"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    return 


def check_hours_header_vs_ext_my_time(header_data, hours):
    """Check if there are hours in ExtMyTime, the total hours match between ExtMyTime and the report, the general
    activities hours are not more than 8%, the general activities hours in ExtMyTime and the report match and if
    the specific hours in ExtMyTime and the report match."""

    # This Error has to be translated to the HOURS class
    """
    except IndexError:
        error_message = f"  Could not find the name '{person_data.name_irs}' in the list of names file and, consequently, the hours " \
                        f"in the ExtMyTime couldn't be checked"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
        return
    """
    if almost_equal(hours.emt_total, 0):
        error_message = "  No hours found in the EXT MY TIME"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
        return
    if not almost_equal(float(header_data.reported_hours), hours.emt_total):
        error_message = f"  The total hours as found in the report ({header_data.reported_hours}) don't match the " \
                        f"EXT MY TIME hours ({hours.emt_total})"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    general_activities_proportion = hours.emt_general / hours.emt_total * 100
    if general_activities_proportion > 8:
        error_message = f"  The General Activities task took {float(hours.emt_general):.2f} hours, which is a " \
                        f"{float(general_activities_proportion):.2f}% of the total: {hours.emt_total} hours"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    return


def check_tasks_hours_report_vs_ext_my_time(header_data, hours):
    """Check if the hours for each task plan is coincident between the report and ExtMyTime"""
    # Check the general activities task
    
    float(hours.report_general_taskplan)
    if not almost_equal(hours.emt_general, hours.report_general):
        error_message = f"  The General Activities task hours in ExtMyTime ({float(hours.emt_general):.2f}) " \
                        f"is not coincident with the ones declared in the report " \
                        f"({float(hours.report_general):.2f})"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    # THIS ERROR HAS TO BE PASSED TO THE HOURS CLASS
    """
    except ValueError:
        error_message = f"  The code of the General Activities Task Plan in the report does not match the valid " \
                        f"format: {hours.report_general_taskplan}. The number of hours in the Task Plan could not be compared " \
                        f"between ExtMyTime and the report."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    """

    for task_code in hours.report_taskplan_dic.keys():
        if task_code not in hours.emt_specific.keys():
            hours.emt_specific[task_code] = 0
        if not almost_equal(hours.emt_specific[task_code], hours.report_taskplan_dic[task_code]):
            error_message = f"  The hours of Specific Task {task_code} in ExtMyTime " \
                            f"({float(hours.emt_specific[task_code]):.2f}) are not coincident with the ones declared " \
                            f"in the report ({float(hours.report_taskplan_dic[task_code]):.2f})"
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    for task_code in hours.emt_specific.keys():
        if task_code not in hours.report_taskplan_dic.keys():
            error_message = f"  The hours of Specific Task {task_code} in ExtMyTime were not declared" \
                            f"in the report."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    """
    except ValueError:
        error_message = f"  The code of the General Activities Task Plan in the report does not match the valid " \
                        f"format: {hours.report_general_taskplan}. The number of hours in the Task Plan could not be compared " \
                        f"between ExtMyTime and the report."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    """
    
    """
    # THIS ERROR HAS TO BE TRANSLATED TO THE HOURS CLASS
    except ValueError:
        error_message = f"  The code of a Specific Task Plan in the report does not match the valid format: " \
                        f"{specific_task}. The number of hours in the Task Plan could not be compared between " \
                        f"ExtMyTime and the report."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    """
    return


def get_codes_activities_section(document, cell_ref):
    """
    Gets the numeric codes of the general task and the specific tasks in a section of the text.

    Arguments:
        document (str): whole text of the document
        start_text (str): string that limits the start of the section
        end_text (str): string that limits the end of the section

    Returns:
        str: code of the general task
        list: with the codes of the specific tasks
    """
    general_taskplan_code = None
    specific_taskplans_codes = []
    # Trim the text to only the wanted part
    section = document.tables[cell_ref["table"]].cell(*cell_ref["cell"]).text

    while True:
        try:
            match = re.search(r"Task\s.+(F4E-OMF-1159|General).+", section)
            section = section[match.span()[1]:]
        except AttributeError:
            break
        line = match.group(0)
        if "General Activities".casefold() in line.casefold():
            general_taskplan_code = line[line.find("(") + 1:line.find(")")]
        else:
            specific_taskplans_codes = specific_taskplans_codes + [line[line.find("(") + 1:line.find(")")]]

    return general_taskplan_code, specific_taskplans_codes


def check_codes_sections(header_data, section, document, cell_ref, hours):
    """
    Checks if the codes of tasks in the text are the same as in the Excel file (Hours Task Plan)

    Parameters:
        header_data (class): data read in the header of the document
        folder (str): Name of folder where 'HoursTaskPlan.xlsx' is located
        general_code (str): code of the general task in the document
        specific_codes (list): list of the codes of the specific tasks in the document
        name (class): configurations of the person's name
        section (str): number of section in the document
    """
    general_code, specific_codes = get_codes_activities_section(document, cell_ref)
    if general_code is None:
        error_message = f"  No General Activity code in section {section} could be found in the report. " \
                        f"Check if the format of the code is correct."
        print(error_message)    
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    elif general_code != hours.report_general_taskplan:
        error_message = f"  In section {section}, the General Activity code '{general_code}' cannot be found in " \
                        f"the Task Plan Hours. Either the format of the code is not correct or the number of the " \
                        f"activity code is not correct."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    if len(specific_codes) == 0:
        error_message = f"  No Specific Activity code in section {section} could be found in the report. " \
                        f"Check if the format of the code is correct."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    for code in specific_codes:
        if not int(code) in hours.emt_specific.keys():
            error_message = f"  In section {section}, the Specific Activity code '{code}' cannot be found in the" \
                            f"Task Plan Hours. Either the format of the code is not correct or the number of the" \
                            f"activity code is not correct."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    # THIS ERROR HAS TO BE PASSED TO THE HOURS CLASS
    """
    if len(hours_person) == 0:
        error_message = f"  The name of the person {person_data}, could not be found in the list with the hours in " \
                        f"ExtMyTime. The correspondence of the codes in section 2.2 and 2.4 with the Task Plan " \
                        f"could not be checked."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    """
    return


def check_filename(filename, header_data):
    """
    Checks if the name of the file follows the structure

    Arguments:
        filename (str): name of the file
    """
    name_split = filename.split()
    if name_split[1] != "Monthly" or name_split[2] != "Report":
        error_message = f"  The name of the file is not according to the template. It should be as follows: " \
                        f"'F4E-OMF-1159-01-01-XX Monthly Report Name Surname #YY MZZ 2023.docx'"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]


def check_dates_section3(document, header_data):

    date_author = document.tables[DATE_AUTHOR["table"]].cell(*DATE_AUTHOR["cell"]).text
    date_approval = document.tables[DATE_APPROVAL["table"]].cell(*DATE_APPROVAL["cell"]).text
   
    if date_author != date_approval:
        error_message = f"  The dates in Section 3 are not the same({date_author} vs {date_approval})."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    # Get month from report number in header
    month_header = re.match(r"#\d+_M(\d+)_\d+", header_data.report_number).group(1)  # e.g. "3"
    month_header = int(month_header)
    # Get year from report number in header
    year_header = re.match(r"#\d+_M\d+_(\d+)", header_data.report_number).group(1)  # eg."2022"
    year_header = int(year_header)

    date_author = date_author.replace("Date: ", "")
    date = date_author.split('/')
    month_section3 = int(date[1])
    day_section3 = int(date[0])
    year_section3 = int(date[2])
    if (month_header+1) % 12 != month_section3:
        error_message = f"  The month in Section 3 ({month_section3}) does not correspond to the following month " \
                        f"of the month being reported ({month_header})."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    if day_section3 > 8:
        error_message = f"  The day in Section 3 ({day_section3}) does not correspond to the first 8 days of the month."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    if year_section3 != year_header:
        if month_header != 12:
            error_message = f"  The year in Section 3 ({year_section3}) does not correspond to the correct date of " \
                            f"signature of the report."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    return

def check_text_forbidden_words(text: str, header_data):
    forbidden = ["F4E Project Manager", "F4E Manager", "F4E Line Manager"]
    for word in forbidden:
        if word.lower() in text.lower():
            error_message = f"  The expression '{word}' appears in the body of the document, please delete it."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return

def forbidden_words(document, header_data):
    sections = [NEW_MILESTONE, CURRENT_MILESTONE, MILESTONE_TO_COPY]
    for section in sections:
        check_text_forbidden_words(document.tables[section["table"]].cell(*section["cell"]).text, header_data)

    return
    

def check_months_header(document, header_data):
    month = header_data.report_number.split('_')[1]

    for period in PERIODS:
        line = document.tables[period["table"]].cell(*period["cell"]).text
        if period["table"] == 4:
            next_month = int(month[1:]) % 12 + 1
            month = f"M{next_month:02}"
        if month not in line:
            error_message = f"  The month in the header of Section {period['section']} is not valid."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    
    return


def check_encryption(mr, header_data):
    dms = mr.tables[DMS_CELL["table"]].cell(*DMS_CELL["cell"]).text
    try:
        token = mr.tables[KEY_ENCRYPTED["table"]].cell(*KEY_ENCRYPTED["cell"]).text
    except IndexError:
        error_message = f"  Could not access the ecnrypted key in Section 3, probably the pre-processing " \
                        f"tool was not used."
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
        return
    token = token.split('\n')[1][2:-1]
    dms_decoded = decode_token(token)
    if dms != dms_decoded:
        error_message = f"  The DMS does not correspond to the encrypted key, the pre-processing tool " \
                        f"was not used."
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


def decode_token(token):
    key = b'XHHzTu2MlETGr1Dy3ltNATwnsuCCaZqgGCp0Dkw0HB4='
    f = Fernet(key)
    dms = f.decrypt(token.encode('utf-8'))
    return dms.decode('utf-8')
    
    
def almost_equal(float_1, float_2):
    return abs(float_1 - float_2) <= 0.0001


def process_monthly(filename, hours_task_plan, list_employees):
    # Read list of employees
    
    name_file = "F4E-OMF-1159-01-01-106 Monthly Report Urszula Pitera #09 M03 2025.docx"

    global name_report
    
    #print(f"Analyzing {filename.name}...")
    #f4e_contract = filename.name.split()[0]
    f4e_contract = name_file.split()[0]
    #name_report = ' '.join(filename.name.split()[3:-3])
    name_report = unidecode(" ".join(name_file.split()[3:-3]))
    #name_report = unidecode(re.match(r".+ Monthly Report (.+\s.+) #", filename).group(1))
    document = docx.Document(filename)
    # Get header fields
    header_data = read_header(document)
    # Check if the name of the file follows correct structure
    #check_filename(filename.name, header_data)
    check_filename(name_file, header_data)
    # Get information from LIST OF EMPLOYEES
    person_data = PersonData(list_employees)
    person_data.select_row(name_report)
    # Get Hours from ExtMyTime and report
    hours = Hours()
    hours.hours_extmytime(hours_task_plan, person_data)
    hours.hours_report(document, header_data)
    # Shows revision number in the header
    show_version_message(header_data)
    # Checks if F4E contracts is the same in the name of the report and the header
    check_f4e_contract(f4e_contract, header_data, person_data)
    # Check if DMS in the header and in "DMS Number Monthly Report.xlsx" are the same
    check_supplier_dms(header_data, person_data)
    # Check if number of report (#) is coherent with months passed from KoM
    check_report_number_against_kom_date(header_data)
    # Check if the F4E reference is the same in header and external file
    check_customer_ref(header_data, person_data)
    # Check if the total number of hours in section 2.3 is the same as in the header
    check_hours_report_vs_header(header_data, hours)
    check_hours_header_vs_ext_my_time(header_data, hours)
    # Check if hours for each task plan is the same in the report and ExtMyTime
    check_tasks_hours_report_vs_ext_my_time(header_data, hours)
    # Check numerical Codes of tasks in sections 2.2 and 2.4
    check_codes_sections(header_data, "2.2", document, NEW_MILESTONE, hours)
    check_codes_sections(header_data, "2.4", document, MILESTONE_TO_COPY, hours)
    # Check both dates in section 3 are the same
    check_dates_section3(document, header_data)
    # Check there are no "forbidden words" in the text
    forbidden_words(document, header_data)
    # Check months headers
    check_months_header(document, header_data)
    # Check encrypted key
    check_encryption(document, header_data)
    
    return


if __name__ == "__main__":
    mr_files = [r"D:\DATA\ferrmar\Documents\04-ATG\automatic_monthly_check\webapp\Development\utils\F4E-OMF-1159-01-01-106 Monthly Report Urszula Pitera #09 M03 2025.docx"]
    hours_task_plan = r"D:\DATA\ferrmar\Documents\04-ATG\automatic_monthly_check\webapp\Development\utils\HoursTaskPlan 2.xlsx"
    process_mr(mr_files, hours_task_plan)
