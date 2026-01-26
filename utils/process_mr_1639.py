import re
from datetime import datetime
import docx
import pandas as pd
from unidecode import unidecode
from cryptography.fernet import Fernet
from xml.etree.cElementTree import XML  

# PARAMETERS
REPORT_NUMBER = {"table": 0, "cell": (0, 2)}
VERSION = {"table": 0, "cell": (0, 4)}
F4E_REFERENCE = {"table": 0, "cell": (1, 2)}
DMS_CELL = {"table": 0, "cell": (1, 4)}
KOM_DATE = {"table": 1, "cell": (1, 1)}
H_IN_PERIOD_CELL = {"table": 1, "cell": (1, 3)}
PERIODS = [
    {"table": 2, "cell": (0, 1), "section": "2.2"},
    {"table": 3, "cell": (0, 1), "section": "2.3"},
    {"table": 4, "cell": (0, 1), "section": "2.4"},
    {"table": 5, "cell": (0, 1), "section": "2.5"}
]
AUTHOR_NAME = {"table": 7, "cell": (1, 0)} 
DATE_AUTHOR = {"table": 7, "cell": (2, 0)} 
DATE_APPROVAL = {"table": 7, "cell": (2, 1)} 
NEW_MILESTONE = {"table": 2, "cell": (1, 1)}
CURRENT_MILESTONE = {"table": 3, "cell": (1, 1)}
MILESTONE_TO_COPY = {"table": 5, "cell": (1, 1)} 
TOTAL_HOURS = {"table": 1, "cell": (1, 3)}
SECTION3 = {"table": 3, "cell": (1, 1)}
TO_HIGHLIGHT = [
    SECTION3,
    {"table": 6, "cell": (1, 1)}, 
    MILESTONE_TO_COPY,
]
KEY_ENCRYPTED = {"table": 7, "cell": (3, 0)}
KEY_ENCRYPTED_SIDE = {"table": 7, "cell": (3, 1)}
HOURS_TABLE = {"table": 4, "cell": (1, 1)}

"""# PATTERNS
month_number_pat = re.compile(r"M\d+")
date_pat = re.compile(r"(\d{2})/(\d{2})/(\d{4})")
period_pat = re.compile(r"\d{2}/\d{2}/\d{4}\s*[-–]\s*\d{2}/\d{2}/\d{4}")
pat_file_name = re.compile(r"#\d+\s+M\d+\s+\d+")
#
pat_task = re.compile(r"Task\s+\d+\s\(\d+\)")
pat_number = re.compile(r"\(\d+\)")
pat_hours = re.compile(r"[\d\.]+\s+hours")

re.search(r"task.+:\s*(\d+([,.]\d+)?)\s*hours?\s*", section.lower())"""

FOLDER = "test_data"
MONTH_NUMBER_TO_NAME = {
    1: "January",
    2: "February",
    3: "March",
    4: "April",
    5: "May",
    6: "June",
    7: "July",
    8: "August",
    9: "September",
    10: "October",
    11: "November",
    12: "December",
}

WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
TEXT = WORD_NAMESPACE + "t"

global hours_task_plan

def process_mr_1639(mr_files, hours_timetell):

    global results_df
    results_df = pd.DataFrame(columns=["Reference", "Name", "Error"])
    #hours_task_plan = pd.read_excel(hours_task_plan, skiprows=3)
    hours_timetell = pd.read_excel(hours_timetell)
    #list_employees = pd.read_excel(r"D:\DATA\ferrmar\Documents\04-ATG\automatic_monthly_check\OMF-1639 version\LIST OF EMPLOYEES 1639.xlsx")
    list_employees = pd.read_excel("LIST OF EMPLOYEES 1639.xlsx")
    #list_employees = list_employees[list_employees["Contract status"] == "Active"]
    for report in mr_files:
        process_monthly(report, list_employees, hours_timetell)

    return results_df


#### FUNCTIONS PART 1 ####
def get_data_from_filename(filename):
    f4e_contract = filename.name.split()[0]
    name_report = unidecode(' '.join(filename.name.split()[3:-3]))
    return f4e_contract, name_report


#### FUNCTIONS PART 2 ####
def accept_all_changes(document):
    cells = [REPORT_NUMBER, VERSION, F4E_REFERENCE, DMS_CELL, KOM_DATE, H_IN_PERIOD_CELL,
             AUTHOR_NAME, DATE_AUTHOR, DATE_APPROVAL, NEW_MILESTONE, CURRENT_MILESTONE, MILESTONE_TO_COPY,
             TOTAL_HOURS, SECTION3, PERIODS[0], PERIODS[1], PERIODS[2]]
    
    for cell in cells:
        cell_text = document.tables[cell["table"]].cell(*cell["cell"])
        new_cell_text = ""
        for para in cell_text.paragraphs:
            new_cell_text += get_accepted_text(para) + "\n"
        document.tables[cell["table"]].cell(*cell["cell"]).text = new_cell_text
    return document


def get_accepted_text(p):
    xml = p._p.xml
    if "w:del" in xml or "w:ins" in xml:
        tree = XML(xml)
        runs = (node.text for node in tree.iter(TEXT) if node.text)
        return "".join(runs)
    else:
        return p.text

#### CLASSES PARTS 3 AND 4 ####
class HeaderData:
    def __init__(self):
        self.report_number = None
        self.f4e_reference = None
        self.supplier_dms = None
        self.kom_date = None
        self.reported_hours = None
        self.version = None
        #self.customer_ref = None

    @property
    def totally_filled(self):
        return (
            self.report_number
            and self.f4e_reference
            and self.supplier_dms
            and self.kom_date
            and self.reported_hours
            and self.version
        )

    def __str__(self):
        text = f"""
Report number: {self.report_number}
F4E reference: {self.f4e_reference}
Supplier DMS: {self.supplier_dms}
KOM Date: {self.kom_date}
Reported hours: {self.reported_hours}
Revision: {self.version}"""
        return text


class PersonData:
    def __init__(self, list_employees):
        self.df = list_employees
        self.contract = None
        self.kom = None
        self.dms = None
        self.name_monthly = None
        self.name_irs = None
        self.name_atg = None
        self.name_timetell = None
        self.row_data = None

    def select_row(self, name_report):
        if name_report == "Raul del Val":
            name_report = "Raul Del Val"
        print(unidecode(name_report))
        self.row_data = self.df[self.df["Employee"].astype(str).apply(unidecode).str.lower() == unidecode(name_report).lower()]
        self.define_data()
        return
    
    def define_data(self):
        self.contract = self.row_data["Specific Contract"].values[0].strip()
        self.kom = self.row_data["Kick-Off Meeting"].values[0]
        self.name_monthly = unidecode(self.row_data["Employee"].values[0])
        self.name_irs = unidecode(self.row_data["Employee"].values[0])
        self.name_atg = unidecode(self.row_data["ATG Account Name"].values[0])
        return
    
    def get_dms(self, month, year):
        self.dms = self.row_data[f"{month} {year}"].values[0]
        return
    

class Hours:
    def __init__(self):
        self.table24_general = None
        self.table24_specific = None
        self.table24_total = None
        self.report23_total = None
        self.report23_general = None
        self.report23_specific = None
        self.report23_taskplan_dic = None
        self.report23_general_taskplan = None
        self.ttexported_general = None
        self.ttexported_specific = None
        self.ndays_worked = None
    
    def hours_timetell_export(self, hours_ttexport, person_data, header_data):
        """Get hours from Excel file exported from TimeTell"""
        self.ttexported_general = 0
        self.ttexported_specific = {}
        
        #format and clean dataframe
        hours_ttexport = self._clean_timetell_df(hours_ttexport)
        if len(hours_ttexport) == 0:
            error_message = f"  The TimeTell export file does not contain any hours for the project F4E-OMF-1639"
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
            return True
        
        #create new column changing name format
        hours_ttexport['name_tt'] = hours_ttexport["Employee name"].str.split(",").str[1].str[1:] + " " + hours_ttexport["Employee name"].str.split(",").str[0]
        print(person_data.name_atg)
        #filter by name
        hours_tt = hours_ttexport[hours_ttexport["name_tt"].astype(str).apply(unidecode) == person_data.name_atg]
        
        #create error if no hours
        if len(hours_tt) == 0:
            error_message = f"  The name '{person_data.name_atg}' could not be found in the TimeTell export file and, consequently, the hours " \
                            f"couldn't be checked"
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
            return True
        
        # Delete word Task and calculate total hours per task. 
        tasks_hours = hours_tt[['Hours', 'Activity name']].groupby('Activity name').sum()
        self.ttexported_specific = tasks_hours.to_dict()['Hours']

        # Days with more than 10 hours worked
        daily_hours = hours_tt[['Hours', 'Date']].groupby('Date').sum()
        days_over_10h = daily_hours[daily_hours['Hours'] > 10]
        self.days_over_10h = days_over_10h.index.tolist()

        # Calculate number of days worked in the period
        self.ndays_worked = hours_tt['Date'].nunique()
        print(self.report23_general_taskplan)
        self.ttexported_general = self.ttexported_specific[self.report23_general_taskplan]
        self.ttexported_specific.pop(self.report23_general_taskplan)

        return False

    def hours_table_section_24(self, document):
        """Get hours from table in section 2.4 of the MR"""
        self.table24_general = 0
        self.table24_specific = {}
        hours_table = document.tables[HOURS_TABLE["table"]].cell(*HOURS_TABLE["cell"]).tables[0]

        # ERROR IN CASE THE TABLE IS NOT FOUND

        for row in hours_table.rows[1:]:
            #Change '.' to ',' in hours and transform to float
            print(row.cells[5].text.strip())
            hours_task = float(row.cells[5].text.strip().replace(",", "."))
            key = str(row.cells[2].text.strip())
            #add line by line the hours in each task. 
            if "Total".casefold() in row.cells[0].text.strip().casefold():
                continue
            if key == self.report23_general_taskplan:
                self.table24_general += hours_task
            elif key in self.table24_specific:
                self.table24_specific[key] += hours_task
            else:
                self.table24_specific[key] = hours_task
        self.table24_total = self.table24_general + sum(self.table24_specific.values())

        return

    def hours_section_23(self, document, header_data):
        """Get hours reported in section 2.3 of the MR"""
        self.report23_general = 0
        self.report23_specific = 0
        self.report23_general_taskplan = ""
        self.report23_taskplan_dic = {}
        i = 0
        section = document.tables[CURRENT_MILESTONE["table"]].cell(*CURRENT_MILESTONE["cell"]).text
        """section = document.tables[CURRENT_MILESTONE["table"]].cell(*CURRENT_MILESTONE["cell"])
        pat_task = re.compile(r"Task\s+\d+\s\(\d+\)")
        pat_number = re.compile(r"\(\d+\)")
        pat_hours = re.compile(r"[\d\.]+\s+hours")
        for par in section.paragraphs:
            # if it finds a task pattern in the paragraph
            if pat_task.search(par.text) is not None:
                # identify the task code
                try:
                    task = pat_number.search(par.text).group().strip("(").strip(")")
                except AttributeError:
                    raise RuntimeError(f"Could not find a task number in {par.text}")
                # get the hours
                try:
                    hours = pat_hours.search(par.text).group().strip(" hours")
                    hours = float(hours)
                    
                except AttributeError:
                    raise RuntimeError(f"Could not find hours in {par.text}")
                # assign hours to task code. 
                if "General Activities".casefold() in par.text.casefold():
                    self.report23_general += hours
                    self.report23_general_taskplan = task
                else:
                    self.report23_specific += hours
                    self.report23_taskplan_dic[task] = hours
        self.report23_total = self.report23_general + self.report23_specific
        return"""
        
        while True and i < 10:
            # look for the line where task is reported
            try:
                match = re.search(r"task.+:\s*(\d+([,.]\d+)?)\s*hours?\s*", section.lower())
                section = section[match.span()[1]:]
            except AttributeError:
                break

            #get hours
            hours_task = match.group(1).replace(",", ".")
            line = match.group(0)

            # assign hours to task code. 
            try:
                hours_task = float(hours_task)
                if "General Activities".casefold() in line.casefold():
                    self.report23_general += hours_task
                    self.report23_general_taskplan = line[line.find("(") + 1:line.find(")")]
                else:
                    self.report23_specific += hours_task
                    self.report23_taskplan_dic[line[line.find("(") + 1:line.find(")")]] = hours_task
            except ValueError:
                error_message = f"  The number of hours in the line '{line}' could not be transformed to a number." \
                                f"Check if it is written correctly. The hours of that task could not be processed and," \
                                f"probably, the total number of hour will be incorrect because of this."
                print(error_message)
                results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
            
        self.report23_total = self.report23_general + self.report23_specific
        i += 1

        return


    def _clean_timetell_df(self, df):
        """Cleans the dataframe imported from TimeTell export"""
        # Drop columns not needed
        df.drop(columns=['Client name', 'Organization name', 'Info', 'Year', 'Month'], inplace=True, errors='ignore')
        #delete row if no name present
        df = df[df['Employee name'].notna()].reset_index(drop=True)
        # filter only activities of the project
        df = df[df['Project name'].notna()]
        mask_activities = df['Project name'].str.contains('F4E-OMF-1639')
        df = df[mask_activities].reset_index()
        # filter to more than 0 hours and nan values
        df = df.dropna(subset=['Hours'])
        df = df[df['Hours'] > 0]
        #round hours to 2 decimals
        df['Hours'] = df['Hours'].astype(float)
        df['Hours'] = df['Hours'].round(2)
        # Conditioning of columns
        df['Activity name'] = df['Activity name'].astype(str)
        df['From time'] = df['From time'].dt.round('min')
        df['To time'] = df['To time'].dt.round('min')
        # Delete word 'Task: ' from activity name
        df['Activity name'] = df['Activity name'].astype(str).str.replace('Task: ', '', regex=False)
        df['Activity name'] = df['Activity name'].astype(float).astype(int).astype(str)
        return df

#### FUNCTIONS PART 3 ####
def read_header(document):
    """
    Extracts the data from the header of the monthly report

    Parameters:
        document (str): string with the text of the report

    Returns
        HeaderData class
    """
    header_data = HeaderData()
    header_data.report_number = document.tables[REPORT_NUMBER["table"]].cell(*REPORT_NUMBER["cell"]).text.strip()
    header_data.version = document.tables[VERSION["table"]].cell(*VERSION["cell"]).text.strip()
    header_data.f4e_reference = document.tables[F4E_REFERENCE["table"]].cell(*F4E_REFERENCE["cell"]).text.strip()
    header_data.supplier_dms = document.tables[DMS_CELL["table"]].cell(*DMS_CELL["cell"]).text.strip()
    header_data.kom_date = document.tables[KOM_DATE["table"]].cell(*KOM_DATE["cell"]).text.strip()
    hours = document.tables[H_IN_PERIOD_CELL["table"]].cell(*H_IN_PERIOD_CELL["cell"]).text.strip()
    to_be_replaced = [" ", "(", "*", ")"]
    for symbol in to_be_replaced:
        hours = hours.replace(symbol, "")
    header_data.reported_hours = hours
    
    if not header_data.totally_filled:
        error_message = f"  The monthly header was not successfully read {header_data}"
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
        raise ValueError(error_message)
    return header_data


#### FUNCTIONS PART 5 ####
def check_filename(filename, header_data):
    """
    Checks if the name of the file follows the structure

    Arguments:
        filename (str): name of the file
    """
    name_split = filename.split()
    if name_split[1] != "Monthly" or name_split[2] != "Report":
        error_message = f"  The name of the file is not according to the template. It should be as follows: " \
                        f"'F4E-OMF-1159-01-01-XX Monthly Report Name Surname #YY MZZ 2023.docx'"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


def show_version_message(header_data):
    """
    Prints a messages with the Revision number of the document
    :param header_data(class)
    :return: None
    """
    print(f"  The revision number in the header is {header_data.version}. Make sure this is correct.")
    return


def check_f4e_contract(code_from_filename, header_data, person_data):
    """Check if F4E reference is the same in the name of the report and inside the report"""
    if header_data.f4e_reference not in code_from_filename:
        error_message = f"  The F4E contract shown in the header ({header_data.f4e_reference}) does not match " \
                        f"the one of the Word filename ({code_from_filename})"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    if header_data.f4e_reference != person_data.contract:
        error_message = f"  The F4E contract shown in the header ({header_data.f4e_reference}) does not match " \
                        f"the one in the LIST OF EMPLOYEES file ({person_data.contract})"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    return


def check_supplier_dms(header_data, person_data):
    """Check if the report of the month, year and person is in the Excel file 'DMS Number Monthly Report.xlsx' and
    also if the DMS number is correspondent"""

    # Get month from report number in header
    month = re.match(r"#\d+_M(\d+)_\d+", header_data.report_number).group(1)  # e.g. "3"
    month = int(month)
    month = MONTH_NUMBER_TO_NAME[month]  # e.g. "March"
    # Get year from report number in header
    year = re.match(r"#\d+_M\d+_(\d+)", header_data.report_number).group(1)  # eg."2022"

    # Search DMS and compare to report
    person_data.get_dms(month, year)
    if not pd.isna(person_data.dms):
        #dms_code = dms.values[0]
        if person_data.dms != header_data.supplier_dms:
            error_message = f"  DMS from database ({person_data.dms}) does not match DMS from " \
                            f"header ({header_data.supplier_dms}). Check the DMS number and the month number in the " \
                            f"report."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    else:
        error_message = f"  The DMS reference could not be found for {month} {year} for {person_data.name_monthly}. " \
                        f"Either the DMS is not in the file or any of these parameters is written incorrectly. " \
                        f"It could not be checked if the DMS number is correct."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


def check_report_number_against_kom_date(header_data):
    """Check if the report number is coherent with the months passed since the KoM"""
    # Get KOM date
    day, month, year = [int(x) for x in header_data.kom_date.split(r"/")]
    kom_date = datetime(year, month, day)
    # Get month, year and report number from header
    number, month, year = header_data.report_number.split("_")
    number, month, year = int(number[1:]), int(month[1:]), int(year)
    report_date = datetime(year, month, 1)
    # Check if difference of months is equal to report number
    months_passed = diff_month(report_date, kom_date) + 1
    if months_passed != number:
        error_message = f"  The F4E report number ({header_data.report_number}) and the KOM date " \
                        f"({header_data.kom_date}) have a difference of {months_passed} months"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


def diff_month(d1, d2):
    """
    Returns the difference of months between two dates

    Arguments:
        d1 (datetime)
        d2 (datetime)
    Returns:
        int
    """
    return (d1.year - d2.year) * 12 + d1.month - d2.month


#### FUNCTIONS SECTION 6 ####

def check_hours_report_vs_header(header_data, hours):
    """Check if the hours declared in the header of the report are equal to the
    sum of the hours declared in section 2.3 of the report"""
    
    # Check if sum of hours in the report is the same as in the header
    if not almost_equal(float(header_data.reported_hours), hours.report23_total):
        error_message = f"  The sum of hours of the tasks ({hours.report23_total}) "\
                        f"declared in section 2.3 is not the same as the one found " \
                        f"in the header ({header_data.reported_hours})"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return 


def other_checks_hours(header_data, hours):
    """Perform other checks on the hours:
    -check if total hours in section 2.4 is bigger than 0
    -check if total hours in header match those in section 2.4
    -check if general activities hours are more than 8%
    -check if average hours is bigger than 8 per day"""

    #Check if total hours section 2.4 is bigger than 0
    if almost_equal(hours.table24_total, 0):
        error_message = "  No hours declared in section 2.4 of the report."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]


    #Check if total hours in header match those in section 2.4
    if not almost_equal(float(header_data.reported_hours), hours.table24_total):
        error_message = f"  The total hours reported in the header ({header_data.reported_hours}) don't match the " \
                        f"total hours reported in section 2.4({hours.table24_total})"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    
    #Check if general activities hours are more than 8%
    general_activities_proportion = hours.table24_general / hours.table24_total * 100
    if general_activities_proportion > 8:
        error_message = f"  The General Activities task {hours.report23_general_taskplan} took {float(hours.table24_general):.2f} hours, which is a " \
                        f"{float(general_activities_proportion):.2f}% of the total: {hours.table24_total} hours"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    #check if average hours is bigger than 8 per day
    if (hours.table24_total / hours.ndays_worked) > 8:
        error_message = f"  The total hours reported in section 2.4 is {hours.table24_total}, " \
                        f"which gives an average of more than 8 hours per day worked."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    # check days with more than 10 hours worked
    if len(hours.days_over_10h) > 0:
        error_message = f"  There are {len(hours.days_over_10h)} days with more than 10 hours worked: "
        for day in hours.days_over_10h:
            error_message += f"{day.strftime('%d/%m/%Y')} "
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    
    return


def check_hours_report_vs_ttexport(header_data, hours):
    """Check if the hours for each task plan is coincident between the table
    in section 2.4 of the report and the hours exported from Timetell"""
    
    # Check hours general task
    if not almost_equal(hours.ttexported_general, hours.table24_general):
        error_message = f"  The General Activities task declared in section 2.4 ({hours.table24_general}) are not " \
                        f" coincident with those declared TimeTell({float(hours.ttexported_general):.2f}) "
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, 
                                           name_report, error_message[2:]]

    # Check hours specific tasks
    for task_code in hours.table24_specific.keys():
        if task_code not in hours.ttexported_specific.keys():
            hours.ttexported_specific[task_code] = 0
        if not almost_equal(hours.table24_specific[task_code], hours.ttexported_specific[task_code]):
            error_message = f"  The hours of Specific Task {task_code} in section 2.4 ({hours.table24_specific[task_code]}) "\
                            f"are not coincident with those declared in TimeTell ({float(hours.ttexported_specific[task_code]):.2f})"
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    return


def check_tasks_hours_report_vs_timetell(header_data, hours):
    """Check if the hours declared in section 2.3 and 2.4 of the report are 
    coincident"""

    # Check the general activities task
    float(hours.report23_general_taskplan)
    if not almost_equal(hours.table24_general, hours.report23_general):
        error_message = f"  The General Activities task {hours.report23_general_taskplan} hours "\
                        f"declared in the table of section 2.4 ({float(hours.table24_general):.2f}) " \
                        f"are not coincident with the ones declared in section 2.3 " \
                        f"({float(hours.report23_general):.2f})"
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    # Check the specific tasks
    for task_code in hours.report23_taskplan_dic.keys():
        if task_code not in hours.table24_specific.keys():
            hours.table24_specific[task_code] = 0
        if not almost_equal(hours.table24_specific[task_code], hours.report23_taskplan_dic[task_code]):
            error_message = f"  The Specific task {task_code} hours declared in the table"\
                            f" of section 2.4 ({float(hours.table24_specific[task_code]):.2f}) "\
                            f"are not coincident with the ones declared in section 2.3 " \
                            f"({float(hours.report23_taskplan_dic[task_code]):.2f})"
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    for task_code in hours.table24_specific.keys():
        if task_code not in hours.report23_taskplan_dic.keys():
            error_message = f"  The hours of Specific Task {task_code} in TimeTell were " \
                            f"declared in section 2.4 ({float(hours.table24_specific[task_code]):.2f}) but not " \
                            f"in section 2.3 of the report."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


#### FUNCTIONS SECTION 7 ####
def check_codes_sections(header_data, section, document, cell_ref, hours):
    """
    Checks if the codes of tasks in the text are the same as in the Excel file (Hours Task Plan)

    Parameters:
        header_data (class): data read in the header of the document
        folder (str): Name of folder where 'HoursTaskPlan.xlsx' is located
        general_code (str): code of the general task in the document
        specific_codes (list): list of the codes of the specific tasks in the document
        name (class): configurations of the person's name
        section (str): number of section in the document
    """
    general_code, specific_codes = get_codes_activities_section(document, cell_ref)
    if general_code is None:
        error_message = f"  No General Activity code in section {section} could be found in the report. " \
                        f"Check if the format of the code is correct."
        print(error_message)    
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    elif general_code != hours.report23_general_taskplan:
        error_message = f"  In section {section}, the General Activity code '{general_code}' cannot be found in " \
                        f"the Task Plan Hours. Either the format of the code is not correct or the number of the " \
                        f"activity code is not correct."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    if len(specific_codes) == 0:
        error_message = f"  No Specific Activity code in section {section} could be found in the report. " \
                        f"Check if the format of the code is correct."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    for code in specific_codes:
        if not code in hours.table24_specific.keys():
            error_message = f"  In section {section}, the Specific Activity code '{code}' cannot be found in the " \
                            f"Task Plan Hours. Either the format of the code is not correct or the number of the" \
                            f"activity code is not correct."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


def get_codes_activities_section(document, cell_ref):
    """
    Gets the numeric codes of the general task and the specific tasks in a section of the text.

    Arguments:
        document (str): whole text of the document
        start_text (str): string that limits the start of the section
        end_text (str): string that limits the end of the section

    Returns:
        str: code of the general task
        list: with the codes of the specific tasks
    """
    general_taskplan_code = None
    specific_taskplans_codes = []
    # Trim the text to only the wanted part
    section = document.tables[cell_ref["table"]].cell(*cell_ref["cell"]).text

    while True:
        try:
            match = re.search(r"Task\s.+(F4E-OMF-1639|General).+", section)
            section = section[match.span()[1]:]
        except AttributeError:
            break
        line = match.group(0)
        if "General Activities".casefold() in line.casefold():
            general_taskplan_code = line[line.find("(") + 1:line.find(")")]
        else:
            specific_taskplans_codes = specific_taskplans_codes + [line[line.find("(") + 1:line.find(")")]]

    return general_taskplan_code, specific_taskplans_codes


def check_dates_section3(document, header_data):

    date_author = document.tables[DATE_AUTHOR["table"]].cell(*DATE_AUTHOR["cell"]).text
    date_approval = document.tables[DATE_APPROVAL["table"]].cell(*DATE_APPROVAL["cell"]).text
   
    if date_author != date_approval:
        error_message = f"  The dates in Section 3 are not the same({date_author} vs {date_approval})."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    # Get month from report number in header
    month_header = re.match(r"#\d+_M(\d+)_\d+", header_data.report_number).group(1)  # e.g. "3"
    month_header = int(month_header)
    # Get year from report number in header
    year_header = re.match(r"#\d+_M\d+_(\d+)", header_data.report_number).group(1)  # eg."2022"
    year_header = int(year_header)

    date_author = date_author.replace("Date: ", "")
    date = date_author.split('/')
    month_section3 = int(date[1])
    day_section3 = int(date[0])
    year_section3 = int(date[2])
    print(month_section3, month_header)
    if (month_header%12) + 1 != month_section3:
        error_message = f"  The month in Section 3 ({month_section3}) does not correspond to the following month " \
                        f"of the month being reported ({month_header})."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    if day_section3 > 8:
        error_message = f"  The day in Section 3 ({day_section3}) does not correspond to the first 8 days of the month."
        print(error_message)
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    if year_section3 != year_header:
        if month_header != 12:
            error_message = f"  The year in Section 3 ({year_section3}) does not correspond to the correct date of " \
                            f"signature of the report."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]

    return


def forbidden_words(document, header_data):
    sections = [NEW_MILESTONE, CURRENT_MILESTONE, MILESTONE_TO_COPY]
    for section in sections:
        check_text_forbidden_words(document.tables[section["table"]].cell(*section["cell"]).text, header_data)
    return


def check_text_forbidden_words(text: str, header_data):
    forbidden = ["F4E Project Manager", "F4E Manager", "F4E Line Manager", "Mindfulness"]
    for word in forbidden:
        if word.lower() in text.lower():
            error_message = f"  The expression '{word}' appears in the body of the document, please delete it."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


def check_months_header(document, header_data):
    month = header_data.report_number.split('_')[1]
    for period in PERIODS:
        line = document.tables[period["table"]].cell(*period["cell"]).text
        if period["table"] == 5:
            next_month = int(month[1:]) % 12 + 1
            month = f"M{next_month:02}"
        if month not in line:
            error_message = f"  The month in the header of Section {period['section']} is not valid."
            print(error_message)
            results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


def check_encryption(mr, header_data):
    dms = mr.tables[DMS_CELL["table"]].cell(*DMS_CELL["cell"]).text.strip()
    version = header_data.version
    try:
        token = mr.tables[KEY_ENCRYPTED["table"]].cell(*KEY_ENCRYPTED["cell"]).text
    except IndexError:
        error_message = f"  Could not access the ecnrypted key in Section 3, probably the pre-processing " \
                        f"tool was not used."
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
        return
    token = token.split('\n')[1][2:-1]
    dms_decoded = decode_token(token)
    text_to_compare = dms + version
    if text_to_compare != dms_decoded:
        error_message = f"  The DMS does not correspond to the encrypted key, the pre-processing tool " \
                        f"was not used."
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, error_message[2:]]
    return


def decode_token(token):
    key = b'XHHzTu2MlETGr1Dy3ltNATwnsuCCaZqgGCp0Dkw0HB4='
    f = Fernet(key)
    dms = f.decrypt(token.encode('utf-8'))
    return dms.decode('utf-8')


#### OTHER GENERAL FUNCTIONS ####

def almost_equal(float_1, float_2):
    return abs(float_1 - float_2) <= 0.0001



#### FLOW FUNCTIONS ####

def get_all_hours(document, header_data, person_data, hours_ttexport):
    hours = Hours()
    hours.hours_section_23(document, header_data)
    hours.hours_table_section_24(document)
    hours_flag = hours.hours_timetell_export(hours_ttexport, person_data, header_data)
    return hours, hours_flag


def header_checks(filename, header_data, f4e_contract, person_data):
    # Check if the name of the file follows correct structure
    check_filename(filename.name, header_data)
    # Shows revision number in the header
    show_version_message(header_data)
    # Checks if F4E contracts is the same in the name of the report and the header
    check_f4e_contract(f4e_contract, header_data, person_data)
    # Check if DMS in the header and in "DMS Number Monthly Report.xlsx" are the same
    check_supplier_dms(header_data, person_data)
    # Check if number of report (#) is coherent with months passed from KoM
    check_report_number_against_kom_date(header_data)
    return


def hours_checks(header_data, hours):
    check_hours_report_vs_header(header_data, hours)
    other_checks_hours(header_data, hours)
    check_hours_report_vs_ttexport(header_data, hours)
    check_tasks_hours_report_vs_timetell(header_data, hours)
    return


def other_checks(document, header_data, hours):
    # Check numerical Codes of tasks in sections 2.2 and 2.5
    check_codes_sections(header_data, "2.2", document, NEW_MILESTONE, hours)
    check_codes_sections(header_data, "2.5", document, MILESTONE_TO_COPY, hours)
    # Check both dates in section 3 are the same
    check_dates_section3(document, header_data)
    # Check there are no "forbidden words" in the text
    forbidden_words(document, header_data)
    # Check months headers
    check_months_header(document, header_data)
    # Check encrypted key
    check_encryption(document, header_data)
    return


def no_errors_message(header_data) -> None:
    if results_df[(results_df["Reference"] == header_data.f4e_reference) & 
                   (results_df["Name"] == name_report)].empty:
        results_df.loc[len(results_df)] = [header_data.f4e_reference, name_report, "Monthly Report processed, no errors found."]
    return


def process_monthly(filename, list_employees, hours_ttexport):
    # Read list of employees
    
    global name_report

    #1 Get data from the filename
    f4e_contract, name_report = get_data_from_filename(filename)
    #2 Open document and accept all changes
    document = docx.Document(filename)
    accept_all_changes(document)
    #3 Get data from header of the report
    header_data = read_header(document)
    person_data = PersonData(list_employees)
    person_data.select_row(name_report)
    #4 Get hours
    hours, hours_flag = get_all_hours(document, header_data, person_data, hours_ttexport)
    #5 Header checks
    header_checks(filename, header_data, f4e_contract, person_data)
    #6 Hours checks
    if not hours_flag:
        hours_checks(header_data, hours)
    #7 Other checks
    other_checks(document, header_data, hours)
    # If no error message, add note saying everything is ok
    no_errors_message(header_data)

    return



